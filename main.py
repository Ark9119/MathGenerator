import random
import datetime
import re

from docx import Document


def generate_num(number_of_numbers, min_number, max_number):
    """Генерирует список случайных чисел"""
    list_of_num = []
    for i in range(number_of_numbers):
        list_of_num.append(str(random.randint(min_number, max_number)))
    return list_of_num


def ansver_for_example(example):
    """
    Вычисляет ответ для математического примера
    Поддерживает: +, -, *, /
    """
    try:
        # Очищаем пример от лишних символов
        clean_example = example.replace('=', '').strip()
        # Разрешаем только цифры, операторы и пробелы
        if not re.match(r'^[\d\s\+\-\*\/\.]+$', clean_example):
            raise ValueError('Недопустимые символы в примере')
        result = eval(clean_example)
        # Округляем до 2 знаков после запятой
        result = round(result, 2)
        return result
    except ZeroDivisionError:
        return 'Ошибка: деление на ноль'
    except Exception as e:
        return f'Ошибка: {e}'


def check_rules_for_example(answer, rules):
    """
    Проверяет ответ на соответствие правилам
    :return: True если все правила пройдены, False иначе
    """
    if not isinstance(answer, (int, float)):
        return False
    for rule in rules:
        if rule == 'not_negative':
            # Ответ не должен быть отрицательным
            if answer < 0:
                return False
        elif rule == 'not_float':
            # Ответ должен быть целым числом
            if answer != int(answer):
                return False
        elif rule == 'not_zero':
            # Ответ не должен быть нулём
            if answer == 0:
                return False
    return True


def generate_example(
    number_of_numbers,
    signs,
    rules,
    max_attempts=1000,
    min_number=0,
    max_number=50
):
    """
    Генерирует пример с проверкой правил
    Перегенерирует пока не найдётся подходящий пример
    """
    for attempt in range(max_attempts):
        # Генерируем числа
        nums = generate_num(number_of_numbers, min_number, max_number)
        # Создаём строку примера
        if len(signs) == 1:
            # Один знак для всех операций
            exampl = f' {signs[0]} '.join(nums)
        else:
            # Разные знаки для разных операций
            exampl = nums[0]
            for i in range(len(nums) - 1):
                sign = random.choice(signs)
                exampl += f' {sign} {nums[i + 1]}'
        # Вычисляем ответ
        answer = ansver_for_example(exampl)
        # Проверяем правила
        if check_rules_for_example(answer, rules):
            print(f'✅ Попыток генерации {attempt + 1}: {exampl} = {answer}')
            return f'{exampl} = ', answer
    # Если не удалось за max_attempts
    print(f'⚠️ Не удалось сгенерировать пример за {max_attempts} попыток')
    return None, None


def generate_all_examples(
    number_of_numbers,
    signs,
    number_of_examples,
    rules,
    min_number=0,
    max_number=50
):
    list_all_examples = []
    list_all_answers = []
    for i in range(number_of_examples):
        print(f'\n--- Генерация примера {i + 1}/{number_of_examples} ---')
        example, answer = generate_example(
            number_of_numbers,
            signs,
            rules,
            min_number=min_number,
            max_number=max_number
        )
        list_all_examples.append(example)
        list_all_answers.append(answer)
    return list_all_examples, list_all_answers


def write_in_docx_file(data, file_name='examples.docx'):
    today = datetime.date.today()
    formated_data = today.strftime('%d.%m.%Y')
    doc = Document()
    doc.add_heading(f'Дата: {formated_data}')
    doc.add_paragraph('\n')
    for i, el in enumerate(data, 1):
        doc.add_paragraph(f'{i}) {el} _____')
    doc.save(file_name)
    print(f'✅ Записано {len(data)} примеров в файл {file_name}')


def write_in_txt_file(data, file_name='examples.txt'):
    with open(file_name, 'w', encoding='utf-8') as f:
        for i, el in enumerate(data, 1):
            f.write(f'{i}) {el} \n')


def main(
    number_of_numbers,  # количество чисел в примере
    znak_operacii,  # используемые в примере знаки
    number_of_examples,  # количество искомых примеров
    rules,  # правила для примеров
    min_number=0,
    max_number=50
):
    # Генерируем примеры и ответы
    examples, answers = generate_all_examples(
        number_of_numbers,
        znak_operacii,
        number_of_examples,
        rules,
        min_number,
        max_number
    )
    # Записываем файлы
    print(examples)
    print(answers)
    write_in_txt_file(examples, file_name='examples.txt')
    write_in_docx_file(examples, file_name='examples.docx')
    write_in_txt_file(answers, file_name='answers.txt')
    write_in_docx_file(answers, file_name='answers.docx')


# Функция-обертка, которую будет вызывать интерфейс
def run_generation(config):
    """
    config: словарь с настройками из интерфейса
    """
    examples, answers = generate_all_examples(
        config['numbers_count'],
        config['signs'],
        config['examples_count'],
        config['rules'],
        config['min_number'],
        config['max_number']
    )
    if not examples:
        return (
            False,
            'Не удалось сгенерировать примеры. Попробуйте изменить условия.'
        )
    write_in_txt_file(examples, file_name='examples.txt')
    write_in_docx_file(examples, file_name='examples.docx')
    write_in_txt_file(answers, file_name='answers.txt')
    write_in_docx_file(answers, file_name='answers.docx')
    return True, f'Успешно создано {len(examples)} примеров!'
